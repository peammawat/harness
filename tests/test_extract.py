"""Document text extraction: txt/md/docx, truncation, and failure handling."""
from __future__ import annotations

from io import BytesIO

from app.agent.extract import extract_document_text


def test_extract_plain_text():
    out = extract_document_text("a.txt", "text/plain", b"hello world")
    assert out == "hello world"


def test_extract_markdown_kept_raw():
    md = b"# Title\n\n- item one\n- item two"
    out = extract_document_text("a.md", "text/markdown", md)
    assert "# Title" in out and "item one" in out


def test_extract_truncates_to_max_chars():
    out = extract_document_text("big.txt", "text/plain", b"x" * 1000, max_chars=50)
    assert len(out) <= 50 + 20  # 50 chars + the truncation marker
    assert "ตัดทอน" in out


def test_extract_empty_returns_placeholder():
    out = extract_document_text("empty.txt", "text/plain", b"")
    assert out == "[ไฟล์ empty.txt ไม่มีข้อความที่อ่านได้]"


def test_extract_bad_pdf_returns_placeholder():
    out = extract_document_text("broken.pdf", "application/pdf", b"not really a pdf")
    assert out == "[ไม่สามารถอ่านไฟล์ broken.pdf]"


def test_extract_docx_round_trip():
    from docx import Document

    doc = Document()
    doc.add_paragraph("Hello from docx")
    doc.add_paragraph("Second line")
    buf = BytesIO()
    doc.save(buf)

    out = extract_document_text(
        "doc.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        buf.getvalue(),
    )
    assert "Hello from docx" in out
    assert "Second line" in out
